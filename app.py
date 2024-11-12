import streamlit as st
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from io import BytesIO

# Title of the application
st.title("ETI Performance Evaluation Form")

# Header information
operator_id = st.text_input("Operator Email")
date = st.date_input("Date")
reviewer_id = st.text_input("Reviewer Email")

# Section for scoring
st.subheader("Positioning of the Patient")

# 1. Positioning of the Patient's Head
score1 = st.radio(
    "1. Positioning of Patient's Head: Did the operator properly tilt the head into the sniffing position?",
    options=["Yes", "No"],
    help="Feedback: The angle of neck flexion should be approximately 35 degrees."
)

# 2. Elevation of the Patient's Head
score2 = st.radio(
    "2. Elevation of Patient's Head: Did the operator properly elevate the head?",
    options=["Yes", "No"],
    help="Feedback: The angle of face extension should be approximately 15 degrees."
)

st.subheader("Insertion of Direct Laryngoscopy Blade")

# 3. Grip of Laryngoscope
score3 = st.radio(
    "3. Grip of Laryngoscope: Did the operator have a proper grip on the laryngoscope?",
    options=["Proper Grip", "Improper Grip"],
    help="Feedback: The laryngoscope should be held in the left hand high enough to avoid obstructing blade entry."
)

# 4. Method to Open Mouth
score4 = st.radio(
    "4. Method to Open Mouth: Did the operator open the mouth by scissoring their finger and thumb?",
    options=["Yes", "No, used another way", "Did not adequately open mouth"],
    help="Feedback: Use opposing pressure on the teeth with thumb and middle finger."
)

# 5. Location of the Blade While Inserting Into the Mouth
score5 = st.radio(
    "5. Location of the Blade While Inserting Into the Mouth",
    options=[
        "Right side, sweep left",
        "Middle, sweep tongue",
        "Other"
    ],
    help="Feedback: Start at the right side to sweep the tongue to the left."
)

# 6. Blade Insertion with Respect to the Vallecula
score6 = st.radio(
    "6. Blade Insertion with Respect to the Vallecula",
    options=["In Vallecula", "Under Vallecula"],
    help="Feedback: Pull blade back into the vallecula."
)

# 7. Force Used while Interacting with Vallecula
score7 = st.radio(
    "7. Force Used while Interacting with Vallecula",
    options=["Appropriate Force", "Excessive Force", "Insufficient Force"],
    help="Feedback: Adjust force direction as needed."
)

# 8. Contact with Teeth During Lifting the Blade
score8 = st.radio(
    "8. Contact with Teeth During Lifting the Blade",
    options=["No Contact", "Hit Teeth, No Damage", "Hit Teeth, Caused Damage"],
    help="Feedback: Blade should be lifted at a 45-degree angle away from the operator."
)

# 9. Order of Events for the Insertion of the Laryngoscope
score9 = st.radio(
    "9. Order of Events for the Insertion of the Laryngoscope",
    options=["Correct Order", "Incorrect Order"],
    help="Feedback: Ensure correct sequence during procedure."
)

st.subheader("Achieving the Optimal Laryngeal View")

# 10. Final Blade Position in the Vallecula When Lifting for Optimal View
score10 = st.radio(
    "10. Final Blade Position in the Vallecula When Lifting for Optimal View",
    options=["Correct Position", "Too Shallow", "Too Deep", "Not in Vallecula"],
    help="Feedback: Adjust blade position as necessary."
)

# 11. Blade Position with Respect to the Oropharynx
score11 = st.radio(
    "11. Blade Position with Respect to the Oropharynx",
    options=["Midline", "Restarted from Right", "Not Midline"],
    help="Feedback: Aim for blade to be in midline."
)

# 12. Lift on Laryngoscope for Proper View
score12 = st.radio(
    "12. Lift on Laryngoscope for Proper View",
    options=["Sufficient Lift", "Insufficient Lift"],
    help="Feedback: Increase lift at a 45-degree angle if necessary."
)

# 13. Quality of the Vocal Cords View
score13 = st.radio(
    "13. Quality of the Vocal Cords View",
    options=["Vocal Cords in View", "Vocal Cords Not in View"],
    help="Feedback: Ensure view of vocal cords before intubating."
)

# 14. Angle of Lift on First Attempt
score14 = st.radio(
    "14. Angle of Lift on First Attempt",
    options=["Appropriate Angle (45°)", "Too Shallow", "Backward onto Teeth"],
    help="Feedback: Adjust handle to keep angle around 45 degrees."
)

# 15. Multiple Blade Insertion Attempts to Achieve Proper View
score15 = st.radio(
    "15. Multiple Blade Insertion Attempts to Achieve Proper View",
    options=["First Attempt", "2-3 Attempts", "4+ Attempts", "No Proper View"],
    help="Feedback: Optimize early attempts for efficiency."
)

st.subheader("Inserting the Endotracheal Tube")

# 16. Number of Contacts of Tube During Insertion
score16 = st.radio(
    "16. Number of Contacts of Tube During Insertion",
    options=["No or Negligible Contact", "Excessive Contact", "Tube Not Inserted"],
    help="Feedback: Minimize contact during insertion."
)

# 17. Multiple Intubation Attempts
score17 = st.radio(
    "17. Multiple Intubation Attempts",
    options=["First Attempt", "One Additional Attempt", "Two or More Attempts", "Intubation Unsuccessful"],
    help="Feedback: Ensure clear view of vocal cords before intubating."
)

st.subheader("Avoiding Injury to the Patient")

# 18. Was Excessive Force Used to Insert the Laryngoscope into the Oropharynx?
score18 = st.radio(
    "18. Was Excessive Force Used to Insert the Laryngoscope into the Oropharynx?",
    options=["No", "Yes"],
    help="Feedback: Reduce approach rate and observe tissues."
)

# 19. Was Excessive Force Used to Insert the ETT into the Oropharynx?
score19 = st.radio(
    "19. Was Excessive Force Used to Insert the ETT into the Oropharynx?",
    options=["No", "Yes", "Tube Not Inserted"],
    help="Feedback: Avoid forcing the ETT into the oropharynx."
)

# 20. Was Excessive Force Used While Interacting with the Vocal Cords?
score20 = st.radio(
    "20. Was Excessive Force Used While Interacting with the Vocal Cords?",
    options=["No", "Yes"],
    help="Feedback: Use smooth approach to avoid damaging vocal cords."
)

# 21. Laryngoscope Manipulation Around Lip(s)
score21 = st.radio(
    "21. Laryngoscope Manipulation Around Lip(s)",
    options=["No Pinching", "Pinched Lips"],
    help="Feedback: Ensure lips are clear from laryngoscope blade."
)

# 22. Laryngoscope and ETT Contact with Tissue and Structures
score22 = st.radio(
    "22. Laryngoscope and ETT Contact with Tissue and Structures",
    options=["Appropriate Contact", "Excessive Contact"],
    help="Feedback: Minimize contact with surrounding tissue."
)

# User input for comments after Question 22
comments22 = st.text_area(
    "Comments",
    "Please enter your comments here...",
    height=150,
    help="Provide feedback or observations regarding the contact of the laryngoscope and ETT with surrounding tissue and structures."
)

# Submit button
if st.button("Submit"):
    # Generate Word document
    doc = Document()
    doc.add_heading('ETI Performance Evaluation Form Summary', 0)

    # Add user input summary to the document
    doc.add_paragraph(f"Operator Email: {operator_id}")
    doc.add_paragraph(f"Reviewer Email: {reviewer_id}")
    doc.add_paragraph(f"Date: {date}")
    
    doc.add_heading("Positioning of the Patient", level=1)
    doc.add_paragraph(f"1. Did the Operator Position the Patient's Head Properly? {score1}")
    doc.add_paragraph(f"2. Did the Operator Properly Elevate or Did Not Need to Elevate the Patient's Head? {score2}")

    doc.add_heading("Insertion of Direct Laryngoscopy Blade", level=1)
    doc.add_paragraph(f"3. The Operator Had a Proper Grip on the Laryngoscope. {score3}")
    doc.add_paragraph(f"4. The Operator Adequately Opened the Mouth. {score4}")
    doc.add_paragraph(f"5. Location of the Blade: {score5}")
    doc.add_paragraph(f"6. Blade Insertion with Respect to the Vallecula: {score6}")
    doc.add_paragraph(f"7. Force Used while Interacting with Vallecula: {score7}")
    doc.add_paragraph(f"8. Contact with Teeth During Lifting the Blade: {score8}")
    doc.add_paragraph(f"9. Order of Events for the Insertion of the Laryngoscope: {score9}")

    
    doc.add_heading("Achieving the Optimal Laryngeal View", level=1)
    doc.add_paragraph(f"10. Final Blade Position in the Vallecula When Lifting for Optimal View: {score10}")
    doc.add_paragraph(f"11. Blade Position with Respect to the Oropharynx: {score11}")
    doc.add_paragraph(f"12. Lift on Laryngoscope for Proper View: {score12}")
    doc.add_paragraph(f"13. Quality of the Vocal Cords View: {score13}")
    doc.add_paragraph(f"14. Angle of Lift on First Attempt: {score14}")
    doc.add_paragraph(f"15. Multiple Blade Insertion Attempts to Achieve Proper View: {score15}")

    doc.add_heading("Inserting the Endotracheal Tube", level=1) 
    doc.add_paragraph(f"16. Number of Contacts of Tube During Insertion: {score16}")
    doc.add_paragraph(f"17. Multiple Intubation Attempts: {score17}")

    doc.add_heading("Avoiding Injury to the Patient", level=1)  
    doc.add_paragraph(f"18. Was Excessive Force Used to Insert the Laryngoscope into the Oropharynx?: {score18}")
    doc.add_paragraph(f"19. Was Excessive Force Used to Insert the ETT into the Oropharynx?: {score19}")
    doc.add_paragraph(f"20. Was Excessive Force Used While Interacting with the Vocal Cords?: {score20}")
    doc.add_paragraph(f"21. Laryngoscope Manipulation Around Lip(s): {score21}")
    doc.add_paragraph(f"22. Laryngoscope and ETT Contact with Tissue and Structures: {score22}")

    doc.add_heading("Comments", level=1)
    doc.add_paragraph(f"Comments: {comments22}")

    # Save the document to a BytesIO stream
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # Email configuration (retrieve credentials from Streamlit secrets)
    sender_email = st.secrets["general"]["email"]
    sender_password = st.secrets["general"]["email_password"]
    recipient_email = operator_id  # Send to the operator's email

    # Prepare the email
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = 'ETI Performance Evaluation Form Summary'

    # Attach the Word document
    part = MIMEBase('application', 'octet-stream')
    part.set_payload(doc_stream.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment; filename="ETI_Evaluation_Summary.docx"')
    msg.attach(part)

    try:
        # Connect to SMTP server (e.g., Gmail)
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, msg.as_string())
        st.success("Evaluation summary sent to the operator.")
    except Exception as e:
        st.error(f"Error sending email: {e}")



